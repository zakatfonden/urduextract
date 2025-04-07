# backend.py (with Merging Restored and New Splitting Function)

import io
import google.generativeai as genai
from docx import Document
from docx.shared import Pt # Import Pt for font size if needed
from docx.enum.text import WD_ALIGN_PARAGRAPH # Required for alignment constant
from docx.oxml.ns import qn # For setting complex script font correctly
from docx.oxml import OxmlElement # Needed for creating font elements if missing
import logging
import os
import streamlit as st
import json
from google.cloud import vision
from docxcompose.composer import Composer # Keep for merging
import copy # Needed for deep copying elements during split

# --- Configure Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(module)s - %(message)s')

# --- START: Runtime Credentials Setup for Streamlit Cloud (Unchanged) ---
# (Keep the existing credentials setup block exactly as it was)
# Define the path for the temporary credentials file within the container's filesystem
CREDENTIALS_FILENAME = "google_credentials.json"
_credentials_configured = False # Flag to track if setup was attempted

if "GOOGLE_CREDENTIALS_JSON" in st.secrets:
    logging.info("Found GOOGLE_CREDENTIALS_JSON in Streamlit Secrets. Setting up credentials file.")
    try:
        # 1. Read from secrets and log its representation
        credentials_json_content_from_secrets = st.secrets["GOOGLE_CREDENTIALS_JSON"]
        logging.info(f"Read {len(credentials_json_content_from_secrets)} characters from secret.")
        # Log first 500 chars using repr() to see hidden characters like \n explicitly
        logging.info(f"REPR of secret content (first 500 chars):\n>>>\n{repr(credentials_json_content_from_secrets[:500])}\n<<<")

        if not credentials_json_content_from_secrets.strip():
                logging.error("GOOGLE_CREDENTIALS_JSON secret is empty.")
                _credentials_configured = False
        else:
            # 2. Write the file with potential cleaning
            file_written_successfully = False
            try:
                # --- START CLEANING ATTEMPT ---
                cleaned_content = credentials_json_content_from_secrets
                try:
                    # Attempt to parse to find the private key and clean it specifically
                    temp_data = json.loads(credentials_json_content_from_secrets)
                    if 'private_key' in temp_data and isinstance(temp_data['private_key'], str):
                        original_pk = temp_data['private_key']
                        cleaned_pk = original_pk.replace('\r\n', '\n').replace('\r', '\n').replace('\\n', '\n')
                        if cleaned_pk != original_pk:
                            logging.warning("Attempted to clean '\\r' or incorrectly escaped '\\n' characters from private_key string.")
                            temp_data['private_key'] = cleaned_pk
                            cleaned_content = json.dumps(temp_data, indent=2)
                        else:
                            cleaned_content = credentials_json_content_from_secrets
                    else:
                        logging.warning("Could not find 'private_key' field (or it's not a string) in parsed secret data for cleaning.")
                        cleaned_content = credentials_json_content_from_secrets
                except json.JSONDecodeError:
                    logging.warning("Initial parse for targeted cleaning failed. Trying global replace on raw string (less safe).")
                    cleaned_content = credentials_json_content_from_secrets.replace('\r\n', '\n').replace('\r', '\n').replace('\\n', '\n')
                # --- END CLEANING ATTEMPT ---

                with open(CREDENTIALS_FILENAME, "w", encoding='utf-8') as f:
                    f.write(cleaned_content)
                logging.info(f"Successfully wrote potentially cleaned credentials to {CREDENTIALS_FILENAME} using UTF-8 encoding.")
                file_written_successfully = True
            except Exception as write_err:
                logging.error(f"CRITICAL Error during file writing (with cleaning attempt): {write_err}", exc_info=True)
                _credentials_configured = False # Ensure flag is false on write error

            # 3. If written, read back immediately and verify/parse
            if file_written_successfully:
                credentials_content_read_back = None
                try:
                    with open(CREDENTIALS_FILENAME, "r", encoding='utf-8') as f:
                        credentials_content_read_back = f.read()
                    logging.info(f"Successfully read back {len(credentials_content_read_back)} characters from {CREDENTIALS_FILENAME}.")
                    logging.info(f"REPR of read-back content (first 500 chars):\n>>>\n{repr(credentials_content_read_back[:500])}\n<<<")

                    # 4. Try parsing the read-back content manually using standard json library
                    try:
                        json.loads(credentials_content_read_back)
                        logging.info("Manual JSON parsing of read-back content SUCCEEDED.")
                        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_FILENAME
                        logging.info(f"GOOGLE_APPLICATION_CREDENTIALS set to point to: {CREDENTIALS_FILENAME}")
                        _credentials_configured = True # Mark configuration as successful ONLY here
                    except json.JSONDecodeError as parse_err:
                        logging.error(f"Manual JSON parsing of read-back content FAILED: {parse_err}", exc_info=True)
                        _credentials_configured = False # Parsing failed
                    except Exception as manual_parse_generic_err:
                        logging.error(f"Unexpected error during manual JSON parsing: {manual_parse_generic_err}", exc_info=True)
                        _credentials_configured = False

                except Exception as read_err:
                    logging.error(f"CRITICAL Error reading back credentials file {CREDENTIALS_FILENAME}: {read_err}", exc_info=True)
                    _credentials_configured = False

    except Exception as e:
        logging.error(f"CRITICAL Error reading secret: {e}", exc_info=True)
        _credentials_configured = False

elif "GOOGLE_APPLICATION_CREDENTIALS" in os.environ:
    logging.info("Using GOOGLE_APPLICATION_CREDENTIALS environment variable set externally.")
    if os.environ.get("GOOGLE_APPLICATION_CREDENTIALS") and os.path.exists(os.environ["GOOGLE_APPLICATION_CREDENTIALS"]):
        logging.info(f"External credentials file found at: {os.environ['GOOGLE_APPLICATION_CREDENTIALS']}")
        _credentials_configured = True
    else:
        logging.error(f"External GOOGLE_APPLICATION_CREDENTIALS path not found or not set: {os.environ.get('GOOGLE_APPLICATION_CREDENTIALS')}")
        _credentials_configured = False

else:
    logging.warning("Vision API Credentials NOT found: Neither GOOGLE_CREDENTIALS_JSON secret nor GOOGLE_APPLICATION_CREDENTIALS env var is set.")
    _credentials_configured = False
# --- END: Runtime Credentials Setup ---


# --- PDF/Image Processing with Google Cloud Vision (Unchanged) ---
def extract_text_from_pdf(pdf_file_obj):
    """
    Extracts text from a PDF file object using Google Cloud Vision API OCR.
    Handles both text-based and image-based PDFs.
    Args:
        pdf_file_obj: A file-like object representing the PDF.
    Returns:
        str: The extracted text, with pages separated by double newlines.
             Returns an empty string "" if no text is found.
             Returns an error string starting with "Error:" if a critical failure occurs.
    """
    global _credentials_configured

    if not _credentials_configured:
        logging.error("Vision API credentials were not configured successfully during startup.")
        return "Error: Vision API authentication failed (Credentials setup failed)."

    credentials_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
    if not credentials_path or not os.path.exists(credentials_path):
        logging.error(f"Credentials check failed just before client init: GOOGLE_APPLICATION_CREDENTIALS path '{credentials_path}' not valid or file doesn't exist.")
        return "Error: Vision API credentials file missing or inaccessible at runtime."

    try:
        logging.info(f"Initializing Google Cloud Vision client using credentials file: {credentials_path}")
        client = vision.ImageAnnotatorClient()
        logging.info("Vision client initialized successfully.")

        pdf_file_obj.seek(0)
        content = pdf_file_obj.read()
        file_size = len(content)
        logging.info(f"Read {file_size} bytes from PDF stream.")

        if not content:
            logging.warning("PDF content is empty.")
            return ""

        mime_type = "application/pdf"
        input_config = vision.InputConfig(content=content, mime_type=mime_type)
        features = [vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)]
        # Explicitly setting language hint for Urdu/Arabic script
        image_context = vision.ImageContext(language_hints=["ur", "ar"]) # Added 'ur'
        request = vision.AnnotateFileRequest(
            input_config=input_config, features=features, image_context=image_context
        )

        logging.info("Sending request to Google Cloud Vision API (batch_annotate_files)...")
        response = client.batch_annotate_files(requests=[request])
        logging.info("Received response from Vision API.")

        if not response.responses:
            logging.error("Vision API returned an empty response list.")
            return "Error: Vision API returned no response."

        first_file_response = response.responses[0]

        if first_file_response.error.message:
            error_message = f"Vision API Error for file: {first_file_response.error.message}"
            logging.error(error_message)
            return f"Error: {error_message}"

        all_extracted_text = []
        if not first_file_response.responses:
            logging.warning("Vision API's AnnotateFileResponse contained an empty inner 'responses' list. No pages processed?")
            return ""

        for page_index, page_response in enumerate(first_file_response.responses):
            if page_response.error.message:
                logging.warning(f"  > Vision API Error for page {page_index + 1}: {page_response.error.message}")
                continue
            if page_response.full_text_annotation:
                page_text = page_response.full_text_annotation.text
                all_extracted_text.append(page_text)
            # else: # Optional logging
            #   logging.info(f"  > Page {page_index + 1} had no full_text_annotation.")

        extracted_text = "\n\n".join(all_extracted_text) # Use double newline as page separator

        if extracted_text:
            logging.info(f"Successfully extracted text from {len(all_extracted_text)} page(s) using Vision API. Total Length: {len(extracted_text)}")
            if not extracted_text.strip():
                logging.warning("Vision API processed pages, but extracted text is empty/whitespace after combining.")
                return ""
            return extracted_text
        else:
            logging.warning("Vision API response received, but no usable full_text_annotation found on any page.")
            return ""

    except Exception as e:
        logging.error(f"CRITICAL Error during Vision API interaction: {e}", exc_info=True)
        # Check for specific authentication errors if possible
        if "Could not automatically determine credentials" in str(e):
             return f"Error: Failed to process PDF with Vision API. Authentication failed: {e}"
        return f"Error: Failed to process PDF with Vision API. Exception: {e}"


# --- Gemini Processing (Unchanged from previous version) ---
def process_text_with_gemini(api_key: str, raw_text: str, rules_prompt: str, model_name: str):
    """
    Processes raw text using the specified Gemini API model based on provided rules.

    Args:
        api_key (str): The Gemini API key.
        raw_text (str): The raw text extracted from the PDF.
        rules_prompt (str): User-defined rules/instructions for Gemini.
        model_name (str): The specific Gemini model ID to use (e.g., "gemini-1.5-pro-latest").

    Returns:
        str: The processed text from Gemini. Returns empty string "" if raw_text is empty.
             Returns an error string starting with "Error:" if a failure occurs.
    """
    if not api_key:
        logging.error("Gemini API key is missing.")
        return "Error: Gemini API key is missing."

    if not raw_text or not raw_text.strip():
        logging.warning("Skipping Gemini call: No raw text provided.")
        return ""

    if not model_name:
        logging.error("Gemini model name is missing.")
        return "Error: Gemini model name not specified."

    try:
        genai.configure(api_key=api_key)
        logging.info(f"Initializing Gemini model: {model_name}")
        model = genai.GenerativeModel(model_name)

        full_prompt = f"""
        **Instructions:**
        {rules_prompt}

        **Text to Process (likely Urdu or Arabic script):**
        ---
        {raw_text}
        ---

        **Output:**
        Return ONLY the processed text according to the instructions. Do not add any introductory phrases like "Here is the processed text:". Ensure proper script formatting (e.g., Urdu, Arabic) and right-to-left presentation where appropriate.
        """

        logging.info(f"Sending request to Gemini model: {model_name}. Text length: {len(raw_text)}")
        # Consider adding safety settings if needed
        # safety_settings = [
        #     {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        #     {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        #     {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        #     {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        # ]
        # response = model.generate_content(full_prompt, safety_settings=safety_settings)
        response = model.generate_content(full_prompt)


        # Error handling for response (check for content blocking or empty response)
        if not response.parts:
            block_reason = None
            safety_ratings = None
            finish_reason = 'UNKNOWN' # Default finish reason

            if hasattr(response, 'prompt_feedback'):
                block_reason = getattr(response.prompt_feedback, 'block_reason', None)
                safety_ratings = getattr(response.prompt_feedback, 'safety_ratings', None)
                # Get finish reason if available (might indicate other issues like length)
                finish_reason = getattr(response.prompt_feedback, 'finish_reason', 'UNKNOWN')


            if block_reason:
                block_reason_msg = f"Content blocked by Gemini safety filters. Reason: {block_reason}"
                logging.error(f"Gemini request ({model_name}) blocked. Reason: {block_reason}. Ratings: {safety_ratings}")
                return f"Error: {block_reason_msg}"
            else:
                # Log other reasons for empty response if not blocked
                logging.warning(f"Gemini ({model_name}) returned no parts (empty response). Finish Reason: {finish_reason}. Safety Ratings: {safety_ratings}")
                # Return empty string, maybe the model genuinely produced nothing or hit another limit
                return ""

        # If parts exist, attempt to get text
        try:
             processed_text = response.text
        except ValueError as ve:
             # Handle cases where response.text might raise ValueError (e.g., function calls expected but not handled)
             logging.error(f"Gemini response ({model_name}) did not contain valid text. Potential function call or unexpected format? Error: {ve}", exc_info=True)
             # Check finish reason again, might provide clues
             finish_reason = getattr(getattr(response, 'prompt_feedback', None), 'finish_reason', 'UNKNOWN')
             return f"Error: Gemini response format issue. Finish Reason: {finish_reason}"


        logging.info(f"Successfully received response from Gemini ({model_name}). Processed text length: {len(processed_text)}")
        return processed_text

    except Exception as e:
        logging.error(f"Error interacting with Gemini API ({model_name}): {e}", exc_info=True)
        # Add more specific error messages if possible (e.g., API key invalid, quota exceeded)
        error_detail = str(e)
        if "API key not valid" in error_detail:
             return "Error: Invalid Gemini API Key."
        elif "Quota" in error_detail:
             return "Error: Gemini API Quota Exceeded."
        return f"Error: Failed to process text with Gemini ({model_name}). Details: {e}"


# --- Create SINGLE Word Document (Sets RTL/Font - Unchanged) ---
def create_word_document(processed_text: str):
    """
    Creates a single Word document (.docx) in memory containing the processed text.
    Sets paragraph alignment to right and text direction to RTL for Arabic/Urdu.
    Uses Arial font for complex scripts. Handles splitting text into paragraphs.

    Args:
        processed_text (str): The text to put into the document.

    Returns:
        io.BytesIO: A BytesIO stream containing the Word document data, or None on critical error.
                    Returns stream with placeholder if processed_text is empty.
    """
    try:
        document = Document()
        # --- Apply Formatting to Normal Style ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial' # Recommended for broad compatibility
        font.rtl = True # Enable RTL for the style's font

        # Set complex script font using oxml
        style_element = style.element
        rpr = style_element.xpath('.//w:rPr')
        if not rpr:
            rpr = OxmlElement('w:rPr')
            style_element.append(rpr)
        else:
            rpr = rpr[0]

        font_name_element = rpr.find(qn('w:rFonts'))
        if font_name_element is None:
            font_name_element = OxmlElement('w:rFonts')
            rpr.append(font_name_element)
        font_name_element.set(qn('w:cs'), 'Arial') # Set Complex Script font

        # Set paragraph format defaults for Normal style
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_format.right_to_left = True
        # --- End Style Formatting ---

        if processed_text and processed_text.strip():
            # Split text into paragraphs based on newlines
            lines = processed_text.strip().split('\n')
            for line in lines:
                if line.strip(): # Avoid adding empty paragraphs
                    # Add paragraph - it inherits style defaults (RTL, font, alignment)
                    paragraph = document.add_paragraph(line.strip())
                    # Redundant explicit setting (safe fallback, but style should handle it)
                    # paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # paragraph.paragraph_format.right_to_left = True
                    # for run in paragraph.runs:
                    #     run.font.rtl = True
                    #     run.font.complex_script = True
        else:
            # Handle empty or whitespace-only content
            empty_msg = "[No text extracted or processed for this file]"
            paragraph = document.add_paragraph(empty_msg)
            paragraph.italic = True # Make it visually distinct
            # Ensure formatting is applied even for the placeholder via the style

        # Save document to a BytesIO stream
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)  # Rewind the stream to the beginning for reading
        logging.info("Successfully created single Word document in memory.")
        return doc_stream

    except Exception as e:
        logging.error(f"Error creating single Word document: {e}", exc_info=True)
        return None # Indicate failure to create the document stream


# --- Merging Function using docxcompose (Restored/Kept) ---
def merge_word_documents(doc_streams_data: list[tuple[str, io.BytesIO]]):
    """
    Merges multiple Word documents (provided as BytesIO streams) into one single document
    by direct concatenation without adding separators. Uses docxcompose library.

    Args:
        doc_streams_data (list[tuple[str, io.BytesIO]]): A list of tuples,
            where each tuple is (original_filename_str, word_doc_bytes_io_obj).

    Returns:
        io.BytesIO: A BytesIO stream containing the final merged Word document,
                    or None on error or if input list is empty.
    """
    if not doc_streams_data:
        logging.warning("No document streams provided for merging.")
        return None

    try:
        # Initialize Composer with the first document
        first_filename, first_stream = doc_streams_data[0]
        if not first_stream or not hasattr(first_stream, 'seek'):
             logging.error(f"Invalid stream provided for the first document: {first_filename}")
             return None
        first_stream.seek(0)
        try:
            master_doc = Document(first_stream)
        except Exception as load_err:
             logging.error(f"Failed to load base document '{first_filename}' with python-docx: {load_err}", exc_info=True)
             return None

        logging.info(f"Loaded base document from '{first_filename}'.")
        composer = Composer(master_doc)
        logging.info(f"Initialized merger with base document.")

        # Append remaining documents
        for i in range(1, len(doc_streams_data)):
            filename, stream = doc_streams_data[i]
            if not stream or not hasattr(stream, 'seek'):
                 logging.warning(f"Skipping invalid stream for document: {filename}")
                 continue # Skip this file
            stream.seek(0)
            logging.info(f"Merging content directly from '{filename}'...")
            try:
                 sub_doc = Document(stream)
                 composer.append(sub_doc)
                 logging.info(f"Successfully appended content from '{filename}'.")
            except Exception as append_err:
                 logging.warning(f"Failed to append document '{filename}': {append_err}. Skipping.", exc_info=True)
                 # Continue with the next file

        # Save the final merged document
        merged_stream = io.BytesIO()
        composer.save(merged_stream)
        merged_stream.seek(0)
        logging.info(f"Successfully merged {len(doc_streams_data)} documents (or fewer if errors occurred) by concatenation.")
        return merged_stream

    except Exception as e:
        logging.error(f"Error merging Word documents using docxcompose: {e}", exc_info=True)
        return None

# --- NEW: Splitting Function ---
def split_word_document(merged_doc_buffer: io.BytesIO, paragraphs_per_split: int = 75):
    """
    Splits a large Word document (from a BytesIO buffer) into smaller documents,
    each containing approximately `paragraphs_per_split` paragraphs (and associated tables).
    Attempts to preserve basic RTL and font settings.

    Args:
        merged_doc_buffer (io.BytesIO): Buffer containing the merged Word document.
        paragraphs_per_split (int): Target number of paragraphs per split file.
                                     This is an approximation for page count.

    Returns:
        list[dict]: A list of dictionaries, where each dict is
                    {'filename': 'split_part_N.docx', 'buffer': BytesIO_obj}
                    Returns an empty list if input is invalid or splitting fails.
    """
    if not merged_doc_buffer or not hasattr(merged_doc_buffer, 'seek'):
        logging.error("Invalid or empty buffer provided for splitting.")
        return []

    merged_doc_buffer.seek(0)
    try:
        source_doc = Document(merged_doc_buffer)
        logging.info(f"Loaded merged document for splitting. Found {len(source_doc.paragraphs)} paragraphs total (approx).")
    except Exception as e:
        logging.error(f"Failed to load merged document from buffer for splitting: {e}", exc_info=True)
        return []

    split_results = []
    current_split_doc = None
    current_paragraph_count = 0
    split_file_index = 0

    # --- Helper to create a new document and apply base styles ---
    def create_new_split_document():
        new_doc = Document()
        # Apply basic RTL/Font settings to the 'Normal' style of the new doc
        try:
            style = new_doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.rtl = True

            style_element = style.element
            rpr = style_element.xpath('.//w:rPr')
            if not rpr: rpr = OxmlElement('w:rPr'); style_element.append(rpr)
            else: rpr = rpr[0]

            font_name_element = rpr.find(qn('w:rFonts'))
            if font_name_element is None: font_name_element = OxmlElement('w:rFonts'); rpr.append(font_name_element)
            font_name_element.set(qn('w:cs'), 'Arial')

            paragraph_format = style.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph_format.right_to_left = True
            logging.info(f"Applied base RTL/Font style to new split document {split_file_index + 1}")
        except Exception as style_err:
            logging.warning(f"Could not apply base styles to new split document {split_file_index + 1}: {style_err}")
        return new_doc
    # --- End Helper ---

    # Iterate through block-level items (paragraphs and tables) in the source document body
    # Using document.element.body directly to access underlying XML elements
    for block in source_doc.element.body:
        if current_split_doc is None:
            current_split_doc = create_new_split_document()

        # Check if the block is a paragraph ('w:p') or a table ('w:tbl')
        is_paragraph = block.tag.endswith('p')
        is_table = block.tag.endswith('tbl')

        if is_paragraph or is_table:
            # Create a deep copy of the element to add to the new document
            # This avoids modifying the source document and potential issues with shared objects
            try:
                 new_element = copy.deepcopy(block)
                 current_split_doc.element.body.append(new_element)
            except Exception as copy_err:
                 logging.warning(f"Failed to copy element (type: {block.tag}) to split doc {split_file_index + 1}: {copy_err}. Skipping element.")
                 continue # Skip this element

            # Increment count only for paragraphs to control splitting
            if is_paragraph:
                current_paragraph_count += 1

            # Check if the split point is reached (based on paragraph count)
            if current_paragraph_count >= paragraphs_per_split:
                # Save the current split document to a buffer
                try:
                    split_buffer = io.BytesIO()
                    current_split_doc.save(split_buffer)
                    split_buffer.seek(0)
                    split_filename = f"split_part_{split_file_index + 1}.docx"
                    split_results.append({'filename': split_filename, 'buffer': split_buffer})
                    logging.info(f"Saved split file: {split_filename} ({current_paragraph_count} paragraphs included)")

                    # Reset for the next split file
                    split_file_index += 1
                    current_split_doc = None # Signal to create a new one
                    current_paragraph_count = 0
                except Exception as save_err:
                    logging.error(f"Failed to save split document part {split_file_index + 1}: {save_err}", exc_info=True)
                    # Reset anyway to avoid adding more to a failed doc
                    current_split_doc = None
                    current_paragraph_count = 0


    # Save any remaining content in the last split document
    if current_split_doc is not None and current_paragraph_count > 0:
        try:
            split_buffer = io.BytesIO()
            current_split_doc.save(split_buffer)
            split_buffer.seek(0)
            split_filename = f"split_part_{split_file_index + 1}.docx"
            split_results.append({'filename': split_filename, 'buffer': split_buffer})
            logging.info(f"Saved final split file: {split_filename} ({current_paragraph_count} paragraphs included)")
        except Exception as save_err:
             logging.error(f"Failed to save final split document part {split_file_index + 1}: {save_err}", exc_info=True)

    logging.info(f"Splitting complete. Generated {len(split_results)} split document(s).")
    return split_results
```


```python
# app.py (Modified for Merge-then-Split Workflow)

import streamlit as st
import backend  # Assumes backend.py is in the same directory
import os
import io
import zipfile # For creating zip files
import logging
import time # For calculating estimates

# Configure basic logging if needed
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Streamlit Page Configuration (Kept from previous request) ---
st.set_page_config(
    page_title="Urdu extraction",
    page_icon="📄",
    layout="wide"
)

# --- Initialize Session State ---
default_state = {
    # 'merged_doc_buffer': None, # No longer stored long-term
    # 'individual_results': [], # Renamed for clarity
    'split_results': [], # NEW: List to store {'filename': 'split_N.docx', 'buffer': BytesIO_obj} from splitting
    'zip_buffer': None, # Buffer for the downloadable zip file of split docs
    'files_processed_count': 0, # Counts original PDFs processed
    'split_files_count': 0, # Counts final split docx files generated
    'processing_complete': False,
    'processing_started': False,
    'ordered_files': [],  # List to hold UploadedFile objects in custom order
}
for key, value in default_state.items():
    if key not in st.session_state:
        st.session_state[key] = value

# --- Helper Functions ---
def reset_processing_state():
    """Resets state related to processing results and status."""
    st.session_state.split_results = [] # NEW
    st.session_state.zip_buffer = None
    st.session_state.files_processed_count = 0
    st.session_state.split_files_count = 0 # NEW
    st.session_state.processing_complete = False
    st.session_state.processing_started = False

# move_file (Unchanged)
def move_file(index, direction):
    files = st.session_state.ordered_files
    if not (0 <= index < len(files)): return
    new_index = index + direction
    if not (0 <= new_index < len(files)): return
    files[index], files[new_index] = files[new_index], files[index]
    st.session_state.ordered_files = files
    reset_processing_state()

# remove_file (Unchanged)
def remove_file(index):
    files = st.session_state.ordered_files
    if 0 <= index < len(files):
        removed_file = files.pop(index)
        st.toast(f"Removed '{removed_file.name}'.")
        st.session_state.ordered_files = files
        reset_processing_state()
    else:
        st.warning(f"Could not remove file at index {index} (already removed or invalid?).")

# handle_uploads (Unchanged)
def handle_uploads():
    if 'pdf_uploader' in st.session_state and st.session_state.pdf_uploader:
        current_filenames = {f.name for f in st.session_state.ordered_files}
        new_files_added_count = 0
        for uploaded_file in st.session_state.pdf_uploader:
            if uploaded_file.name not in current_filenames:
                st.session_state.ordered_files.append(uploaded_file)
                current_filenames.add(uploaded_file.name)
                new_files_added_count += 1
        if new_files_added_count > 0:
            st.toast(f"Added {new_files_added_count} new file(s) to the end of the list.")
            reset_processing_state()

# clear_all_files_callback (Unchanged)
def clear_all_files_callback():
    st.session_state.ordered_files = []
    if 'pdf_uploader' in st.session_state:
        st.session_state.pdf_uploader = []
    reset_processing_state()
    st.toast("Removed all files from the list.")

# --- Create Zip Buffer Helper (Unchanged from previous request) ---
def create_zip_buffer(results_list):
    """Creates a zip file in memory containing multiple docx files."""
    if not results_list:
        return None
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for item in results_list:
            filename = item.get('filename')
            buffer = item.get('buffer')
            if filename and buffer and hasattr(buffer, 'getvalue'):
                buffer.seek(0)
                zipf.writestr(filename, buffer.getvalue())
            else:
                logging.warning(f"Skipping item in zip creation due to missing filename or buffer: {item.get('original_pdf_name', 'Unknown')}") # original_pdf_name won't exist here
    zip_buffer.seek(0)
    return zip_buffer

# --- Page Title & Description (Kept from previous request) ---
st.title("📄 Urdu extraction - PDF to Word Converter")
st.markdown("Upload PDF files (Urdu or Arabic script recommended), arrange order, process, merge, split into ~10-page parts, and download as a ZIP archive.") # Updated description

# --- Sidebar (Kept from previous request - API Key, Model Select, Rules) ---
st.sidebar.header("⚙️ Configuration")

# API Key Input
api_key_from_secrets = st.secrets.get("GEMINI_API_KEY", "")
api_key = st.sidebar.text_input(
    "Enter your Google Gemini API Key", type="password",
    help="Required. Get your key from Google AI Studio.", value=api_key_from_secrets or ""
)
# API Key Status Messages
if api_key_from_secrets and api_key == api_key_from_secrets: st.sidebar.success("API Key loaded from Secrets.", icon="✅")
elif not api_key_from_secrets and not api_key: st.sidebar.warning("API Key not found or entered.", icon="🔑")
elif api_key and not api_key_from_secrets: st.sidebar.info("Using manually entered API Key.", icon="⌨️")
elif api_key and api_key_from_secrets and api_key != api_key_from_secrets: st.sidebar.info("Using manually entered API Key (overrides secret).", icon="⌨️")

# Model Selection
st.sidebar.markdown("---")
st.sidebar.header("🧠 AI Model")
model_options = {
    "Gemini 1.5 Flash (Fastest, Cost-Effective)": "gemini-1.5-flash-latest",
    "Gemini 1.5 Pro (Advanced, Slower, Higher Cost)": "gemini-1.5-pro-latest",
}
pro_model_key = "Gemini 1.5 Pro (Advanced, Slower, Higher Cost)"
pro_model_index = list(model_options.keys()).index(pro_model_key) if pro_model_key in model_options else 0
selected_model_display_name = st.sidebar.selectbox(
    "Choose the Gemini model for processing:",
    options=list(model_options.keys()),
    index=pro_model_index, # Default Pro
    key="gemini_model_select",
    help="Select the AI model. Pro is more capable but slower and costs more."
)
selected_model_id = model_options[selected_model_display_name]
st.sidebar.caption(f"Selected model ID: `{selected_model_id}`")

# Extraction Rules
st.sidebar.markdown("---")
st.sidebar.header("📜 Extraction Rules")
# Default Rules Text (Kept from previous request)
default_rules = """Remove footnotes.
Identify and Completely Remove the header: Find the entire original top line of the page. This usually includes a page number and a title/heading (like كتاب الزكاة ), also content that may exist above the top line. All of this must be removed.
Do not remove headings inside main body text.
Structure the text into logical paragraphs based on the original document. Don't translate anything."""
rules_prompt = st.sidebar.text_area(
    "Enter the rules Gemini should follow:", value=default_rules, height=250,
    help="Provide clear instructions for how Gemini should process the extracted text."
)
# --- End Sidebar ---


# --- Main Area ---
st.header("📁 Manage Files for Processing")

# File Uploader (Unchanged)
uploaded_files_widget = st.file_uploader(
    "Choose PDF files to add to the list below:", type="pdf", accept_multiple_files=True,
    key="pdf_uploader",
    on_change=handle_uploads,
    label_visibility="visible"
)

st.markdown("---")

# --- TOP: Buttons Area & Progress Indicators ---
st.subheader("🚀 Actions & Progress (Top)")
col_b1_top, col_b2_top = st.columns([3, 2])

with col_b1_top:
    # Label reflects the overall process now
    process_button_top_clicked = st.button(
        "✨ Process, Merge & Split Files (Top)", # CHANGED Label
        key="process_button_top",
        use_container_width=True, type="primary",
        disabled=st.session_state.processing_started or not st.session_state.ordered_files
    )

with col_b2_top:
    # Download button for the ZIP of SPLIT files
    if st.session_state.zip_buffer and not st.session_state.processing_started:
        st.download_button(
            label=f"📥 Download All ({st.session_state.split_files_count}) Split Files (.zip)", # CHANGED Label
            data=st.session_state.zip_buffer,
            file_name="split_urdu_documents.zip", # CHANGED Filename
            mime="application/zip",
            key="download_zip_button_top",
            use_container_width=True
        )
    elif st.session_state.processing_started:
        st.info("Processing in progress...", icon="⏳")
    else:
        st.markdown("*(Download button for ZIP of split files appears here)*") # CHANGED Placeholder Text

# Placeholders for top progress indicators
progress_bar_placeholder_top = st.empty()
status_text_placeholder_top = st.empty()

st.markdown("---") # Separator before file list

# --- Interactive File List (Unchanged) ---
st.subheader(f"Files in Processing Order ({len(st.session_state.ordered_files)}):")
if not st.session_state.ordered_files:
    st.info("Use the uploader above to add files. They will appear here for ordering.")
else:
    # Header/Rows (Unchanged)
    col_h1, col_h2, col_h3, col_h4, col_h5 = st.columns([0.5, 5, 1, 1, 1])
    with col_h1: st.markdown("**#**"); with col_h2: st.markdown("**Filename**")
    with col_h3: st.markdown("**Up**"); with col_h4: st.markdown("**Down**"); with col_h5: st.markdown("**Remove**")
    for i, file in enumerate(st.session_state.ordered_files):
        col1, col2, col3, col4, col5 = st.columns([0.5, 5, 1, 1, 1])
        with col1: st.write(f"{i+1}"); with col2: st.write(file.name)
        with col3: st.button("⬆️", key=f"up_{i}", on_click=move_file, args=(i, -1), disabled=(i == 0), help="Move Up")
        with col4: st.button("⬇️", key=f"down_{i}", on_click=move_file, args=(i, 1), disabled=(i == len(st.session_state.ordered_files) - 1), help="Move Down")
        with col5: st.button("❌", key=f"del_{i}", on_click=remove_file, args=(i,), help="Remove")
    st.button("🗑️ Remove All Files", key="remove_all_button", on_click=clear_all_files_callback, help="Click to remove all files from the list.", type="secondary")

st.markdown("---") # Separator after file list

# --- BOTTOM: Buttons Area & Progress Indicators ---
st.subheader("🚀 Actions & Progress (Bottom)")
col_b1_bottom, col_b2_bottom = st.columns([3, 2])

with col_b1_bottom:
    process_button_bottom_clicked = st.button(
        "✨ Process, Merge & Split Files (Bottom)", # CHANGED Label
        key="process_button_bottom",
        use_container_width=True, type="primary",
        disabled=st.session_state.processing_started or not st.session_state.ordered_files
    )

with col_b2_bottom:
    # Download button for the ZIP of SPLIT files
    if st.session_state.zip_buffer and not st.session_state.processing_started:
        st.download_button(
            label=f"📥 Download All ({st.session_state.split_files_count}) Split Files (.zip)", # CHANGED Label
            data=st.session_state.zip_buffer,
            file_name="split_urdu_documents.zip", # CHANGED Filename
            mime="application/zip",
            key="download_zip_button_bottom",
            use_container_width=True
        )
    elif st.session_state.processing_started:
        st.info("Processing in progress...", icon="⏳")
    else:
        st.markdown("*(Download button for ZIP of split files appears here)*") # CHANGED Placeholder Text

# Placeholders for bottom progress indicators
progress_bar_placeholder_bottom = st.empty()
status_text_placeholder_bottom = st.empty()

# --- Container for Individual File Results (Displayed below bottom progress) ---
results_container = st.container()


# --- Processing Logic ---
if process_button_top_clicked or process_button_bottom_clicked:
    reset_processing_state()
    st.session_state.processing_started = True

    # Re-check conditions (Unchanged checks)
    if not st.session_state.ordered_files:
        st.warning("⚠️ No files in the list to process.")
        st.session_state.processing_started = False
    elif not api_key:
        st.error("❌ Please enter or configure your Gemini API Key in the sidebar.")
        st.session_state.processing_started = False
    elif not rules_prompt:
        st.warning("⚠️ The 'Extraction Rules' field is empty. Processing with default model behavior might be less predictable.")
    elif not selected_model_id:
        st.error("❌ No Gemini model selected in the sidebar.")
        st.session_state.processing_started = False

    # Proceed only if checks passed
    if st.session_state.ordered_files and api_key and st.session_state.processing_started and selected_model_id:

        # Store intermediate doc streams for merging
        intermediate_doc_streams = [] # List of tuples: (original_filename, BytesIO_buffer)

        total_files = len(st.session_state.ordered_files)
        TIME_PER_PDF_ESTIMATE_S = 200 # Estimate in seconds (kept from previous)

        # Initialize progress bars
        # Add extra steps for merge/split in progress calculation if desired, or keep simple file-based progress
        total_steps = total_files + 2 # Add 2 steps for merge and split
        current_step = 0

        progress_bar_top = progress_bar_placeholder_top.progress(0, text="Starting processing...")
        progress_bar_bottom = progress_bar_placeholder_bottom.progress(0, text="Starting processing...")

        # Clear previous results visually
        results_container.empty()

        # --- Stage 1: Process each file individually ---
        files_processed_ok_count = 0
        for i, file_to_process in enumerate(st.session_state.ordered_files):
            current_step = i + 1
            progress_value = current_step / total_steps

            original_filename = file_to_process.name
            current_file_status = f"'{original_filename}' ({i + 1}/{total_files})"

            # Calculate Estimated Remaining Time (for file processing part)
            remaining_files_in_stage = total_files - i
            # Adjust estimate if needed, maybe add fixed time for merge/split?
            remaining_time_estimate_s = remaining_files_in_stage * TIME_PER_PDF_ESTIMATE_S + 60 # Add ~1 min for merge/split estimate
            remaining_minutes = int(remaining_time_estimate_s // 60)
            remaining_seconds_part = int(remaining_time_estimate_s % 60)
            time_estimate_str = f"Est. time remaining: {remaining_minutes}m {remaining_seconds_part}s"

            progress_text = f"Processing {current_file_status}. {time_estimate_str}"

            # Update progress bars and status texts
            progress_bar_top.progress(progress_value, text=progress_text)
            progress_bar_bottom.progress(progress_value, text=progress_text)
            status_text_placeholder_top.info(f"🔄 Starting {current_file_status}")
            status_text_placeholder_bottom.info(f"🔄 Starting {current_file_status}")

            with results_container:
                st.markdown(f"--- \n**Processing: {original_filename}**")

            raw_text = None; processed_text = ""; word_doc_stream = None
            extraction_error = False; gemini_error_occurred = False; word_creation_error_occurred = False

            # 1. Extract Text (Error handling as before)
            status_text_placeholder_top.info(f"📄 Extracting text from {current_file_status}...")
            status_text_placeholder_bottom.info(f"📄 Extracting text from {current_file_status}...")
            try:
                file_to_process.seek(0)
                raw_text = backend.extract_text_from_pdf(file_to_process)
                if raw_text is None: raise ValueError("Backend extraction returned None") # Treat None as error
                if isinstance(raw_text, str) and raw_text.startswith("Error:"):
                    with results_container: st.error(f"❌ Error extracting text: {raw_text}")
                    extraction_error = True
                elif not raw_text or not raw_text.strip():
                    with results_container: st.warning(f"⚠️ No text extracted. Placeholder Word content will be used.")
                    processed_text = ""
            except Exception as ext_exc:
                with results_container: st.error(f"❌ Text extraction failed: {ext_exc}")
                extraction_error = True

            # 2. Process with Gemini (if text extracted)
            if not extraction_error and raw_text and raw_text.strip():
                status_text_placeholder_top.info(f"🤖 Sending text to Gemini ({selected_model_display_name})...")
                status_text_placeholder_bottom.info(f"🤖 Sending text to Gemini ({selected_model_display_name})...")
                try:
                    processed_text_result = backend.process_text_with_gemini(api_key, raw_text, rules_prompt, selected_model_id)
                    if processed_text_result is None: raise ValueError("Backend Gemini processing returned None")
                    if isinstance(processed_text_result, str) and processed_text_result.startswith("Error:"):
                        with results_container: st.error(f"❌ Gemini processing error: {processed_text_result}")
                        gemini_error_occurred = True; processed_text = ""
                    else:
                        processed_text = processed_text_result
                        if not processed_text.strip():
                             with results_container: st.warning(f"⚠️ Gemini returned empty text.")
                except Exception as gem_exc:
                    with results_container: st.error(f"❌ Gemini processing failed: {gem_exc}")
                    gemini_error_occurred = True; processed_text = ""

            # 3. Create Intermediate Word Document (always attempt if extraction didn't fail critically)
            if not extraction_error:
                status_text_placeholder_top.info(f"📝 Creating intermediate Word doc for {current_file_status}...")
                status_text_placeholder_bottom.info(f"📝 Creating intermediate Word doc for {current_file_status}...")
                try:
                    word_doc_stream = backend.create_word_document(processed_text) # Use potentially empty processed_text
                    if word_doc_stream:
                        intermediate_doc_streams.append((original_filename, word_doc_stream))
                        files_processed_ok_count += 1 # Count files that resulted in a stream for merging
                        with results_container:
                             success_msg = f"✅ Created intermediate Word file for merging."
                             if gemini_error_occurred: success_msg += " (Used placeholder due to Gemini error)"
                             elif not processed_text and raw_text and raw_text.strip(): success_msg += " (Used placeholder as Gemini result was empty)"
                             elif not processed_text and (not raw_text or not raw_text.strip()): success_msg += " (Based on empty extracted text)"
                             st.success(success_msg)
                    else:
                        word_creation_error_occurred = True
                        with results_container: st.error(f"❌ Failed to create intermediate Word file (backend returned None).")
                except Exception as doc_exc:
                    word_creation_error_occurred = True
                    with results_container: st.error(f"❌ Error creating intermediate Word file: {doc_exc}")
            else:
                 with results_container: st.warning(f"ℹ️ Skipped intermediate Word file creation due to text extraction failure.")

            # Update status text (optional - could show summary here)

        # --- End of file processing loop ---
        st.session_state.files_processed_count = files_processed_ok_count

        # --- Stage 2: Merge Documents ---
        merged_doc_buffer = None
        if intermediate_doc_streams:
            current_step = total_files + 1
            progress_value = current_step / total_steps
            merge_status_text = f"Merging {len(intermediate_doc_streams)} intermediate document(s)..."
            progress_bar_top.progress(progress_value, text=merge_status_text)
            progress_bar_bottom.progress(progress_value, text=merge_status_text)
            status_text_placeholder_top.info(f"💾 {merge_status_text}")
            status_text_placeholder_bottom.info(f"💾 {merge_status_text}")

            with results_container: st.markdown("---"); st.info(f"💾 {merge_status_text}")

            try:
                merged_doc_buffer = backend.merge_word_documents(intermediate_doc_streams)
                if not merged_doc_buffer:
                    with results_container: st.error("❌ Document merging failed (backend returned None). Cannot proceed to splitting.")
                else:
                     with results_container: st.success("✅ Intermediate documents merged successfully.")
            except Exception as merge_exc:
                with results_container: st.error(f"❌ Document merging failed: {merge_exc}. Cannot proceed to splitting.")
                merged_doc_buffer = None # Ensure it's None on error
        else:
            with results_container: st.warning("⚠️ No intermediate documents were created successfully. Skipping merge and split.")

        # --- Stage 3: Split Merged Document ---
        split_results_final = []
        if merged_doc_buffer:
            current_step = total_files + 2
            progress_value = current_step / total_steps
            split_status_text = "Splitting merged document into parts (approx. 10 pages each)..."
            progress_bar_top.progress(progress_value, text=split_status_text)
            progress_bar_bottom.progress(progress_value, text=split_status_text)
            status_text_placeholder_top.info(f"✂️ {split_status_text}")
            status_text_placeholder_bottom.info(f"✂️ {split_status_text}")

            with results_container: st.markdown("---"); st.info(f"✂️ {split_status_text}")

            try:
                # Use the new backend function - adjust paragraphs_per_split if needed
                split_results_final = backend.split_word_document(merged_doc_buffer, paragraphs_per_split=75)
                st.session_state.split_results = split_results_final # Store in session state
                st.session_state.split_files_count = len(split_results_final)

                if not split_results_final:
                     with results_container: st.warning("⚠️ Splitting resulted in zero output files (check merged content or splitting logic).")
                else:
                     with results_container: st.success(f"✅ Merged document split into {st.session_state.split_files_count} part(s).")

            except Exception as split_exc:
                with results_container: st.error(f"❌ Document splitting failed: {split_exc}")
                split_results_final = [] # Ensure empty on error
        else:
             with results_container: st.info("ℹ️ Skipping splitting because merging failed or produced no output.")


        # --- Stage 4: Create Zip Archive ---
        final_status_message = ""
        rerun_needed = False
        if split_results_final:
            status_text_placeholder_top.info("📦 Creating final ZIP archive...")
            status_text_placeholder_bottom.info("📦 Creating final ZIP archive...")
            with results_container: st.info("📦 Creating final ZIP archive...")
            try:
                zip_buffer_final = create_zip_buffer(split_results_final)
                if zip_buffer_final:
                    st.session_state.zip_buffer = zip_buffer_final # Store for download
                    final_status_message = f"✅ Processing complete! Generated {st.session_state.split_files_count} split Word file(s). Click 'Download All' above or below."
                    with results_container: st.success(final_status_message)
                    rerun_needed = True # Rerun to show download buttons
                else:
                    final_status_message = "❌ Failed to create final ZIP archive."
                    with results_container: st.error(final_status_message)
            except Exception as zip_exc:
                final_status_message = f"❌ Error during ZIP archive creation: {zip_exc}"
                with results_container: st.error(final_status_message)
        else:
            # Determine final message if splitting didn't happen or failed
            if not intermediate_doc_streams:
                 final_status_message = "⚠️ Processing finished, but no documents were successfully processed or merged."
            elif not merged_doc_buffer:
                 final_status_message = "⚠️ Processing finished, but merging failed. No split files generated."
            else: # Merging likely worked, but splitting failed or produced nothing
                 final_status_message = "⚠️ Processing finished, merging succeeded, but splitting failed or produced no files."
            with results_container: st.warning(final_status_message)


        # --- Final Cleanup ---
        progress_bar_placeholder_top.empty()
        status_text_placeholder_top.empty()
        progress_bar_placeholder_bottom.empty()
        status_text_placeholder_bottom.empty()

        st.session_state.processing_complete = True
        st.session_state.processing_started = False

        if rerun_needed:
            st.rerun() # Rerun to make download buttons visible

    else:
        # Processing didn't start due to initial checks failing
        st.session_state.processing_started = False # Ensure reset


# --- Fallback info message (Unchanged) ---
if not st.session_state.ordered_files and not st.session_state.processing_started and not st.session_state.processing_complete:
    st.info("Upload PDF files using the 'Choose PDF files' button above.")

# --- Footer (Unchanged) ---
st.markdown("---")
st.markdown("Developed with Streamlit, Google Gemini, and Google Cloud Vision.")
```

**Explanation of Changes:**

* **`backend.py`:**
    * The `merge_word_documents` function is kept (or restored if you removed it).
    * A new function `split_word_document` is added.
        * It takes the merged document buffer and a `paragraphs_per_split` argument (defaulting to 75 as an estimate for ~10 pages).
        * It loads the source document.
        * It iterates through paragraphs and tables (`source_doc.element.body`).
        * It creates new `Document` objects for each split part and applies basic RTL/Arial styling to their 'Normal' style to maintain formatting.
        * It uses `copy.deepcopy()` to copy elements (paragraphs/tables) to the new document to avoid issues.
        * It saves each completed split document into a `BytesIO` buffer.
        * It returns a list of dictionaries: `[{'filename': 'split_part_N.docx', 'buffer': buffer}, ...]`.
* **`app.py`:**
    * The core processing logic inside the `if process_button_top_clicked or process_button_bottom_clicked:` block is restructured:
        1.  **Individual Processing Loop:** Still iterates through uploaded files, performs extraction, Gemini processing, and creates *intermediate* Word documents using `backend.create_word_document`. These are stored in a temporary list `intermediate_doc_streams`.
        2.  **Merging Step:** After the loop, if `intermediate_doc_streams` is not empty, it calls `backend.merge_word_documents` to get a single `merged_doc_buffer`.
        3.  **Splitting Step:** If merging was successful, it calls the new `backend.split_word_document` with the `merged_doc_buffer`. The results (list of split file dicts) are stored in `st.session_state.split_results`.
        4.  **Zipping Step:** If splitting produced results, it calls `create_zip_buffer` with `st.session_state.split_results`. The final zip buffer is stored in `st.session_state.zip_buffer`.
    * **Session State:** Adjusted to store `split_results` and `split_files_count` instead of `individual_results`.
    * **UI:** Button labels, status messages, progress bar text, and download button logic are updated to reflect the merge-then-split workflow and the final zip archive of *split* files.
    * **Progress Bar:** Updated slightly to account for the extra merge and split steps in the total step count for a more accurate representation.
    * **Error Handling:** Added checks and messages for failures during the merge and split stages.

**To Use:**

1.  Replace your `backend.py` with the code from the first block (`backend_py_split`).
2.  Replace your `app.py` with the code from the second block (`app_py_split`).
3.  Ensure you have the necessary libraries: `streamlit`, `google-generativeai`, `python-docx`, `google-cloud-vision`, `docxcompose`. (`copy` and `zipfile` are standard Python libraries).
4.  Run `streamlit run app.py`.

The application will now process files, merge them internally, split the merged result based on paragraph count, and offer a zip file containing these split parts for downlo
